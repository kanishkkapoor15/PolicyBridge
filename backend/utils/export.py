"""Export utilities for generating final Word documents and JSON exports."""

import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def generate_word_export(state: dict) -> bytes:
    """Generate a .docx compliance report with all converted documents and audit trail."""
    doc = Document()

    final_report = state.get("final_report", {})
    converted_docs = state.get("converted_documents", [])
    outstanding = state.get("outstanding_issues", [])
    agent_messages = state.get("agent_messages", [])

    # ── Cover Page ──
    title = doc.add_heading("PolicyBridge Compliance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta_items = [
        f"Session ID: {state.get('session_id', 'N/A')}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Source Jurisdiction: {state.get('source_jurisdiction', 'N/A')}",
        f"Target Jurisdiction: Ireland / European Union",
        f"Policy Category: {state.get('policy_category', 'N/A')}",
        f"Documents Processed: {final_report.get('total_documents_processed', 0)}",
        f"Compliance Score: {final_report.get('compliance_score_before', 0)}% → {final_report.get('compliance_score_after', 0)}%",
    ]
    for item in meta_items:
        p = doc.add_paragraph(item)
        p.style = "List Bullet"

    doc.add_page_break()

    # ── Executive Summary ──
    _add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(final_report.get("executive_summary", "No summary available."))

    # Changes by framework
    changes_by_fw = final_report.get("changes_by_framework", {})
    if changes_by_fw:
        _add_heading(doc, "Changes by Legal Framework", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Framework"
        hdr[1].text = "Changes"
        for fw, count in sorted(changes_by_fw.items(), key=lambda x: -x[1]):
            row = table.add_row().cells
            row[0].text = fw
            row[1].text = str(count)

    # Next steps
    next_steps = final_report.get("next_steps", [])
    if next_steps:
        _add_heading(doc, "Recommended Next Steps", level=2)
        for step in next_steps:
            p = doc.add_paragraph(step)
            p.style = "List Number"

    doc.add_page_break()

    # ── Per-Document Sections ──
    for conv_doc in converted_docs:
        _add_heading(doc, f"Document: {conv_doc.get('filename', 'Unknown')}", level=1)
        doc.add_paragraph(f"Total changes: {conv_doc.get('total_changes', 0)} | Solicitor flags: {len(conv_doc.get('flagged_for_solicitor', []))}")

        for change in conv_doc.get("changes", []):
            _add_heading(doc, f"Change {change.get('change_id', '?')}", level=3)

            # Two-column table: original vs new
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Headers
            table.rows[0].cells[0].text = "Original Clause"
            table.rows[0].cells[1].text = "Converted Clause"
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Content
            orig_cell = table.rows[1].cells[0]
            orig_para = orig_cell.paragraphs[0]
            orig_run = orig_para.add_run(change.get("original_clause", "")[:500])
            orig_run.font.color.rgb = RGBColor(128, 128, 128)  # Grey

            new_cell = table.rows[1].cells[1]
            new_para = new_cell.paragraphs[0]
            new_run = new_para.add_run(change.get("new_clause", "")[:500])
            new_run.font.color.rgb = RGBColor(0, 0, 180)  # Blue

            # Citation
            doc.add_paragraph(f"Legal basis: {change.get('legal_citation', 'N/A')}", style="Intense Quote")
            doc.add_paragraph(f"Reason: {change.get('plain_english_reason', 'N/A')}")
            doc.add_paragraph(f"Confidence: {change.get('confidence', 0)}% | Type: {change.get('change_type', 'N/A')}")

            if change.get("change_type") == "Flagged":
                _add_paragraph(doc, "[SOLICITOR REVIEW REQUIRED]", bold=True, color=RGBColor(200, 0, 0))

            doc.add_paragraph("")

        doc.add_page_break()

    # ── Outstanding Issues Appendix ──
    _add_heading(doc, "Appendix A: Outstanding Issues", level=1)
    if outstanding:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Document"
        hdr[1].text = "Issue"
        hdr[2].text = "Reason for Solicitor Review"
        for issue in outstanding:
            row = table.add_row().cells
            row[0].text = issue.get("filename", issue.get("doc_id", ""))
            row[1].text = issue.get("issue", "")[:200]
            row[2].text = issue.get("reason", "")[:200]
    else:
        doc.add_paragraph("No outstanding issues requiring solicitor review.")

    doc.add_page_break()

    # ── Audit Trail Appendix ──
    _add_heading(doc, "Appendix B: Audit Trail", level=1)
    for msg in agent_messages:
        p = doc.add_paragraph(msg)
        p.style = "List Bullet"
        for run in p.runs:
            run.font.size = Pt(8)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_json_export(state: dict) -> str:
    """Export full session state as formatted JSON."""
    exportable = {}
    for key, value in state.items():
        try:
            json.dumps(value)
            exportable[key] = value
        except (TypeError, ValueError):
            exportable[key] = str(value)
    return json.dumps(exportable, indent=2, default=str)
